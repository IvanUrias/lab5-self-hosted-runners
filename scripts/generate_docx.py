# -*- coding: utf-8 -*-
"""
Script para generar el documento Word del Laboratorio 5
Self-hosted runners y troubleshooting
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LAB_DIR = os.path.dirname(SCRIPT_DIR)
SCREENSHOTS_DIR = os.path.join(LAB_DIR, "docs", "screenshots")
OUTPUT_FILE = os.path.join(LAB_DIR, "Laboratorio 5 — Self-hosted runners y troubleshooting.docx")

# Screenshot mapping
SCREENSHOTS = {
    "repo_created": "github_repo_created.png",
    "settings_runners": "github_settings_runners.png",
    "runner_config": "runner_config_terminal.png",
    "runner_online": "runner_online_github.png",
    "runner_running": "runner_running_terminal.png",
    "hybrid_success": "hybrid_pipeline_success.png",
    "error_label": "troubleshooting_label_error.png",
    "error_dependency": "troubleshooting_dependency_error.png",
    "error_path": "troubleshooting_path_error.png",
    "error_permisos": "troubleshooting_permisos_error.png",
    "fixed_success": "troubleshooting_fixed_success.png",
    "logging_verbose": "logging_verbose_output.png",
}


def setup_styles(doc):
    """Configura estilos personalizados para el documento."""
    # Estilo de título principal
    style = doc.styles['Title']
    font = style.font
    font.size = Pt(28)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    font.bold = True

    # Heading 1
    style = doc.styles['Heading 1']
    font = style.font
    font.size = Pt(18)
    font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    font.bold = True

    # Heading 2
    style = doc.styles['Heading 2']
    font = style.font
    font.size = Pt(14)
    font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    font.bold = True

    # Heading 3
    style = doc.styles['Heading 3']
    font = style.font
    font.size = Pt(12)
    font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    font.bold = True

    # Normal
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)
    font.name = 'Calibri'
    
    return doc


def add_image_with_caption(doc, image_key, caption, width=Inches(5.5)):
    """Añade una imagen con su descripción debajo."""
    image_path = os.path.join(SCREENSHOTS_DIR, SCREENSHOTS.get(image_key, ""))
    
    if os.path.exists(image_path):
        # Añadir imagen centrada
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        
        # Añadir caption
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_p.add_run(f"Figura: {caption}")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # Explicación debajo
        doc.add_paragraph()
    else:
        p = doc.add_paragraph(f"[Captura: {caption} - Imagen no disponible: {image_path}]")
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def add_code_block(doc, code, language="yaml"):
    """Añade un bloque de código formateado."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    # Set background shading
    shading = p.paragraph_format
    p_element = p._element
    pPr = p_element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    pPr.append(shd)


def add_note_box(doc, text, note_type="info"):
    """Añade un cuadro de nota/advertencia."""
    colors = {
        "info": RGBColor(0x0D, 0x47, 0xA1),
        "warning": RGBColor(0xE6, 0x51, 0x00),
        "success": RGBColor(0x1B, 0x5E, 0x20),
        "error": RGBColor(0xB7, 0x1C, 0x1C),
    }
    icons = {
        "info": "ℹ️",
        "warning": "⚠️",
        "success": "✅",
        "error": "❌",
    }
    p = doc.add_paragraph()
    run = p.add_run(f"{icons.get(note_type, '')} {text}")
    run.font.color.rgb = colors.get(note_type, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt(10)


def build_document():
    """Construye el documento completo del laboratorio."""
    doc = Document()
    doc = setup_styles(doc)

    # ==========================================
    # PORTADA
    # ==========================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Laboratorio 5")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Self-hosted runners y troubleshooting")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("GitHub Actions — CI/CD")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_paragraph()
    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Autor: Lucas Arranz del Río")
    run.font.size = Pt(12)

    email = doc.add_paragraph()
    email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = email.add_run("lucas.arranzdelrio@tajamar365.com")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("Mayo 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_page_break()

    # ==========================================
    # ÍNDICE
    # ==========================================
    doc.add_heading("Índice", level=1)
    
    toc_items = [
        "1. Objetivos del laboratorio",
        "2. Escenario y requisitos",
        "3. Parte 1 — Configuración del Runner Self-hosted",
        "   3.1 Creación del repositorio en GitHub",
        "   3.2 Descarga y configuración del runner",
        "   3.3 Registro del runner con labels personalizados",
        "   3.4 Inicio del runner",
        "   3.5 Verificación del runner en GitHub",
        "4. Parte 2 — Workflow Híbrido",
        "   4.1 Diseño del pipeline",
        "   4.2 Código del workflow",
        "   4.3 Ejecución y resultados",
        "5. Parte 3 — Troubleshooting",
        "   5.1 Error: Label incorrecto",
        "   5.2 Error: Dependencia inexistente",
        "   5.3 Error: Path inválido",
        "   5.4 Error: Permisos",
        "   5.5 Versiones corregidas",
        "6. Parte 4 — Logging Detallado",
        "   6.1 Configuración del logging verbose",
        "   6.2 Análisis de los logs",
        "   6.3 Información del runner",
        "7. Parte 5 — Seguridad",
        "   7.1 Riesgos de runners persistentes",
        "   7.2 Riesgos de acceso a red",
        "   7.3 Riesgos de aislamiento insuficiente",
        "8. Conclusiones",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        if item.startswith("   "):
            p.paragraph_format.left_indent = Cm(1.5)
            p.runs[0].font.size = Pt(10)
        else:
            p.runs[0].font.bold = True

    doc.add_page_break()

    # ==========================================
    # 1. OBJETIVOS
    # ==========================================
    doc.add_heading("1. Objetivos del Laboratorio", level=1)
    doc.add_paragraph(
        "Este laboratorio tiene como objetivo practicar los siguientes conceptos "
        "relacionados con GitHub Actions y la infraestructura de CI/CD:"
    )
    objectives = [
        "Self-hosted runners: configuración, registro y gestión de runners en infraestructura propia",
        "Labels: etiquetado de runners para dirigir la ejecución de jobs específicos",
        "Troubleshooting: resolución de errores comunes en pipelines de CI/CD",
        "Logs: análisis de logging detallado para diagnóstico de problemas",
        "Gestión de runners: monitorización del estado y salud de los runners",
        "Diagnósticos: uso de herramientas y variables de entorno para depuración avanzada",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_paragraph()

    # ==========================================
    # 2. ESCENARIO Y REQUISITOS
    # ==========================================
    doc.add_heading("2. Escenario y Requisitos", level=1)
    doc.add_paragraph(
        "La empresa necesita ejecutar pipelines en infraestructura propia. El alumno debe "
        "configurar un runner self-hosted, etiquetarlo correctamente, ejecutar jobs específicos "
        "y resolver errores introducidos deliberadamente."
    )

    doc.add_heading("Entorno de trabajo", level=2)
    env_items = [
        "Sistema Operativo: Windows 11 Enterprise LTSC (Build 26100)",
        "Tipo de sistema: x64-based PC (VMware VM)",
        "Herramientas disponibles: Git 2.x, curl, PowerShell 5.1",
        "Repositorio: https://github.com/lucas-arranz/lab5-self-hosted-runners",
        "Runner: GitHub Actions Runner v2.323.0 para Windows x64",
    ]
    for item in env_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Requisitos del laboratorio", level=2)
    requirements = [
        ("Parte 1 — Runner", "Registrar un runner asociado al repositorio con labels personalizados"),
        ("Parte 2 — Workflow híbrido", "Crear pipeline que ejecute jobs en hosted y self-hosted"),
        ("Parte 3 — Troubleshooting", "Resolver problemas provocados: label incorrecto, dependencia inexistente, path inválido, error de permisos"),
        ("Parte 4 — Logging", "Activar logging detallado e identificar puntos de fallo"),
        ("Parte 5 — Seguridad", "Analizar riesgos de runners persistentes, acceso a red y aislamiento"),
    ]
    for title, desc in requirements:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==========================================
    # 3. PARTE 1 — CONFIGURACIÓN DEL RUNNER
    # ==========================================
    doc.add_heading("3. Parte 1 — Configuración del Runner Self-hosted", level=1)
    doc.add_paragraph(
        "En esta sección se describe el proceso completo de configuración de un "
        "runner self-hosted en GitHub Actions, desde la creación del repositorio "
        "hasta la verificación de que el runner está online y listo para ejecutar jobs."
    )

    # 3.1
    doc.add_heading("3.1 Creación del repositorio en GitHub", level=2)
    doc.add_paragraph(
        "El primer paso es crear un repositorio en GitHub donde se alojarán los workflows "
        "y al cual se asociará el runner self-hosted. Se ha creado el repositorio "
        "lucas-arranz/lab5-self-hosted-runners mediante la API de GitHub."
    )
    doc.add_paragraph(
        "Comando utilizado para crear el repositorio mediante la API REST de GitHub:"
    )
    add_code_block(doc, 
        'curl -X POST -H "Accept: application/vnd.github+json" \\\n'
        '  -H "Authorization: Bearer <TOKEN>" \\\n'
        '  "https://api.github.com/user/repos" \\\n'
        '  -d \'{"name":"lab5-self-hosted-runners","description":"Laboratorio 5","private":false}\''
    )
    doc.add_paragraph(
        "El repositorio se ha inicializado con la siguiente estructura de directorios:"
    )
    add_code_block(doc,
        "LAB 5/\n"
        "├── .github/\n"
        "│   └── workflows/\n"
        "│       ├── hybrid-pipeline.yml\n"
        "│       ├── troubleshooting-errors.yml\n"
        "│       ├── troubleshooting-fixed.yml\n"
        "│       └── logging-detallado.yml\n"
        "├── scripts/\n"
        "│   └── health-check.ps1\n"
        "├── docs/\n"
        "├── .gitignore\n"
        "└── README.md"
    )

    add_image_with_caption(doc, "repo_created",
        "Repositorio lucas-arranz/lab5-self-hosted-runners creado en GitHub con todos los "
        "archivos del proyecto. Se observa la estructura de directorios con los workflows "
        "de GitHub Actions, scripts de diagnóstico y la documentación del proyecto."
    )

    # 3.2
    doc.add_heading("3.2 Descarga y configuración del runner", level=2)
    doc.add_paragraph(
        "Para configurar el runner self-hosted, se descarga el paquete oficial de GitHub "
        "Actions Runner para Windows x64 (v2.323.0). El proceso incluye:"
    )
    steps_32 = [
        "Crear el directorio C:\\actions-runner para alojar el runner",
        "Descargar el archivo ZIP del runner desde GitHub Releases",
        "Extraer el contenido del archivo ZIP",
        "Ejecutar el script de configuración config.cmd",
    ]
    for step in steps_32:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph("Comandos ejecutados:")
    add_code_block(doc,
        "# Crear directorio\n"
        "New-Item -ItemType Directory -Path 'C:\\actions-runner' -Force\n\n"
        "# Descargar runner\n"
        'Invoke-WebRequest -Uri "https://github.com/actions/runner/releases/download/v2.323.0/actions-runner-win-x64-2.323.0.zip" -OutFile "C:\\actions-runner\\actions-runner.zip"\n\n'
        "# Extraer\n"
        'Expand-Archive -Path "C:\\actions-runner\\actions-runner.zip" -DestinationPath "C:\\actions-runner" -Force'
    )

    # 3.3
    doc.add_heading("3.3 Registro del runner con labels personalizados", level=2)
    doc.add_paragraph(
        "El runner se registra en el repositorio utilizando un token de registro temporal "
        "obtenido mediante la API de GitHub. Se asignan los labels personalizados: "
        "self-hosted, windows y lab5."
    )
    doc.add_paragraph("Obtención del token de registro:")
    add_code_block(doc,
        'curl -X POST -H "Accept: application/vnd.github+json" \\\n'
        '  -H "Authorization: Bearer <TOKEN>" \\\n'
        '  "https://api.github.com/repos/lucas-arranz/lab5-self-hosted-runners/actions/runners/registration-token"'
    )
    doc.add_paragraph("Configuración del runner:")
    add_code_block(doc,
        '.\\config.cmd --url https://github.com/lucas-arranz/lab5-self-hosted-runners \\\n'
        '  --token AVYF4EMYDYILVSLOLL6DGDLKCQOSW \\\n'
        '  --name LAB5-Runner \\\n'
        '  --labels self-hosted,windows,lab5 \\\n'
        '  --unattended'
    )

    add_image_with_caption(doc, "runner_config",
        "Terminal de PowerShell mostrando la configuración del runner self-hosted. "
        "Se ejecuta el comando config.cmd con los parámetros de URL del repositorio, "
        "token de registro, nombre del runner (LAB5-Runner) y los labels personalizados "
        "(self-hosted, windows, lab5). El mensaje 'Settings Saved.' confirma que la "
        "configuración se ha guardado correctamente."
    )

    add_note_box(doc, 
        "Los labels asignados (self-hosted, windows, lab5) son fundamentales para "
        "dirigir la ejecución de jobs específicos al runner correcto. El label 'lab5' "
        "es un label personalizado que permite identificar este runner específicamente "
        "para las prácticas de este laboratorio.",
        "info"
    )

    # 3.4
    doc.add_heading("3.4 Inicio del runner", level=2)
    doc.add_paragraph(
        "Una vez configurado, el runner se inicia ejecutando el script run.cmd. "
        "Este proceso mantiene el runner activo y escuchando por nuevos jobs:"
    )
    add_code_block(doc, ".\\run.cmd")
    
    add_image_with_caption(doc, "runner_running",
        "Terminal de PowerShell mostrando el runner self-hosted en ejecución. "
        "Se observa el mensaje '√ Connected to GitHub' indicando la conexión exitosa, "
        "la versión del runner (2.323.0) y el estado 'Listening for Jobs', lo que "
        "confirma que el runner está activo y esperando recibir trabajos del repositorio."
    )

    # 3.5
    doc.add_heading("3.5 Verificación del runner en GitHub", level=2)
    doc.add_paragraph(
        "Tras iniciar el runner, se verifica su estado en la configuración del "
        "repositorio (Settings → Actions → Runners). El runner debe aparecer con "
        "estado 'Idle' (verde), indicando que está online y disponible."
    )

    add_image_with_caption(doc, "settings_runners",
        "Página de configuración de runners del repositorio en GitHub (Settings → Actions → Runners) "
        "antes de registrar el runner. Se muestra el botón 'New self-hosted runner' y el mensaje "
        "indicando que aún no hay runners self-hosted configurados."
    )

    add_image_with_caption(doc, "runner_online",
        "Página de configuración de runners del repositorio mostrando el runner "
        "LAB5-Runner con estado 'Idle' (indicador verde). Se observan los labels "
        "asignados: self-hosted, windows, lab5 y X64. El estado verde confirma que "
        "el runner está online, conectado correctamente y listo para ejecutar jobs."
    )

    doc.add_page_break()

    # ==========================================
    # 4. PARTE 2 — WORKFLOW HÍBRIDO
    # ==========================================
    doc.add_heading("4. Parte 2 — Workflow Híbrido", level=1)
    doc.add_paragraph(
        "Esta sección describe la creación y ejecución de un pipeline híbrido que "
        "combina la ejecución de jobs tanto en runners hosted (proporcionados por GitHub) "
        "como en el runner self-hosted configurado en la parte anterior."
    )

    # 4.1
    doc.add_heading("4.1 Diseño del pipeline", level=2)
    doc.add_paragraph(
        "El pipeline híbrido se ha diseñado con 3 jobs que se ejecutan secuencialmente:"
    )
    
    jobs_desc = [
        ("Job 1 — Build (GitHub Hosted)", 
         "Se ejecuta en ubuntu-latest (runner hosted por GitHub). Realiza el checkout del código, "
         "valida la estructura del proyecto y simula un proceso de build. Genera un artefacto "
         "con información de la compilación."),
        ("Job 2 — Deploy (Self-Hosted)", 
         "Se ejecuta en el runner self-hosted con labels [self-hosted, windows, lab5]. "
         "Descarga el artefacto del job anterior, verifica el entorno local de infraestructura "
         "propia y simula un despliegue. Este job depende del éxito del Job 1."),
        ("Job 3 — Resumen (GitHub Hosted)", 
         "Se ejecuta en ubuntu-latest. Genera un resumen del pipeline completo con el estado "
         "de todos los jobs anteriores. Se ejecuta siempre (if: always()) para reportar incluso "
         "si algún job falla."),
    ]
    for title, desc in jobs_desc:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        p.add_run(desc)
        doc.add_paragraph()

    # 4.2
    doc.add_heading("4.2 Código del workflow", level=2)
    doc.add_paragraph(
        "El archivo hybrid-pipeline.yml define el pipeline completo. A continuación "
        "se muestra la configuración principal de cada job:"
    )

    add_code_block(doc,
        'name: "Lab5 - Pipeline Híbrido"\n\n'
        'on:\n'
        '  push:\n'
        '    branches: [master, main]\n'
        '  workflow_dispatch:\n\n'
        'jobs:\n'
        '  build-hosted:\n'
        '    name: "🏗️ Build (GitHub Hosted)"\n'
        '    runs-on: ubuntu-latest\n'
        '    steps:\n'
        '      - uses: actions/checkout@v4\n'
        '      - name: Validar estructura del proyecto\n'
        '        run: |\n'
        '          test -d ".github/workflows" && echo "✅ OK"\n'
        '      - name: Simular build\n'
        '        run: echo "build_status=success" >> build_info.txt\n\n'
        '  deploy-selfhosted:\n'
        '    name: "🖥️ Deploy (Self-Hosted)"\n'
        '    runs-on: [self-hosted, windows, lab5]\n'
        '    needs: build-hosted\n'
        '    steps:\n'
        '      - uses: actions/checkout@v4\n'
        '      - name: Verificar entorno local\n'
        '        run: systeminfo | Select-String "OS Name"\n'
        '        shell: pwsh\n\n'
        '  summary-hosted:\n'
        '    name: "📊 Resumen (GitHub Hosted)"\n'
        '    runs-on: ubuntu-latest\n'
        '    needs: [build-hosted, deploy-selfhosted]\n'
        '    if: always()'
    )

    add_note_box(doc,
        "La clave del pipeline híbrido es la directiva runs-on. Los jobs hosted usan "
        "'ubuntu-latest' mientras que el job self-hosted usa '[self-hosted, windows, lab5]' "
        "para asegurar que se ejecuta en nuestro runner local.",
        "info"
    )

    # 4.3
    doc.add_heading("4.3 Ejecución y resultados", level=2)
    doc.add_paragraph(
        "Al hacer push del código al repositorio, el workflow se ejecuta automáticamente. "
        "El pipeline completa los 3 jobs exitosamente, demostrando la capacidad de "
        "combinar runners hosted y self-hosted en un mismo pipeline."
    )

    add_image_with_caption(doc, "hybrid_success",
        "Ejecución exitosa del pipeline híbrido en GitHub Actions. Los tres jobs se muestran "
        "con marcas verdes de éxito: 'Build (GitHub Hosted)' ejecutado en ubuntu-latest, "
        "'Deploy (Self-Hosted)' ejecutado en el runner local LAB5-Runner, y 'Resumen "
        "(GitHub Hosted)' ejecutado nuevamente en ubuntu-latest. El flujo secuencial "
        "demuestra la correcta dependencia entre jobs y la diferenciación real de runners."
    )

    doc.add_page_break()

    # ==========================================
    # 5. PARTE 3 — TROUBLESHOOTING
    # ==========================================
    doc.add_heading("5. Parte 3 — Troubleshooting", level=1)
    doc.add_paragraph(
        "En esta sección se documentan los errores deliberadamente introducidos en los "
        "workflows y el proceso de resolución de cada uno. El objetivo es practicar la "
        "identificación y corrección de problemas comunes en GitHub Actions."
    )

    # 5.1
    doc.add_heading("5.1 Error: Label incorrecto", level=2)
    doc.add_heading("Descripción del error", level=3)
    doc.add_paragraph(
        "Se configura un job con runs-on: [self-hosted, nonexistent-label-xyz], "
        "un label que no corresponde a ningún runner registrado. Esto provoca que "
        "el job se quede en estado 'Queued' indefinidamente esperando un runner que "
        "nunca aparecerá."
    )
    doc.add_paragraph("Código del error:")
    add_code_block(doc,
        'error-label-incorrecto:\n'
        '  name: "❌ Error: Label Incorrecto"\n'
        '  runs-on: [self-hosted, nonexistent-label-xyz]  # ← Label que no existe\n'
        '  steps:\n'
        '    - name: Este paso nunca se ejecutará\n'
        '      run: echo "No debería verse"'
    )

    add_image_with_caption(doc, "error_label",
        "Job 'Error: Label Incorrecto' en estado 'Queued' (círculo amarillo). GitHub Actions "
        "muestra el mensaje de advertencia: 'This job is waiting for a runner to be online "
        "that matches the following labels: self-hosted, nonexistent-label-xyz'. El job nunca "
        "se ejecutará porque no existe ningún runner con el label 'nonexistent-label-xyz'."
    )

    doc.add_heading("Resolución", level=3)
    doc.add_paragraph(
        "La solución consiste en corregir el label para que coincida con los labels "
        "realmente asignados al runner (self-hosted, windows, lab5):"
    )
    add_code_block(doc,
        'fix-label-corregido:\n'
        '  name: "✅ Fix: Label Corregido"\n'
        '  runs-on: [self-hosted, windows, lab5]  # ← Labels correctos\n'
        '  steps:\n'
        '    - name: Verificar runner correcto\n'
        '      run: echo "Runner: ${{ runner.name }}"'
    )
    add_note_box(doc,
        "Lección aprendida: Siempre verificar que los labels del runs-on coincidan "
        "exactamente con los labels configurados en el runner. Un solo label incorrecto "
        "impedirá la ejecución del job.",
        "warning"
    )

    # 5.2
    doc.add_heading("5.2 Error: Dependencia inexistente", level=2)
    doc.add_heading("Descripción del error", level=3)
    doc.add_paragraph(
        "Se referencia una GitHub Action que no existe en el marketplace: "
        "actions/this-action-does-not-exist@v99. Esto produce un error inmediato "
        "al intentar resolver la acción."
    )
    add_code_block(doc,
        'error-dependencia-inexistente:\n'
        '  runs-on: ubuntu-latest\n'
        '  steps:\n'
        '    - name: Intentar usar action inexistente\n'
        '      uses: actions/this-action-does-not-exist@v99  # ← No existe'
    )

    add_image_with_caption(doc, "error_dependency",
        "Job 'Error: Dependencia Inexistente' fallido con marca roja. El error muestra: "
        "'Unable to resolve action actions/this-action-does-not-exist@v99, unable to find "
        "version v99'. GitHub Actions no puede descargar ni resolver la acción referenciada "
        "porque simplemente no existe en el registro de actions."
    )

    doc.add_heading("Resolución", level=3)
    doc.add_paragraph(
        "La solución es usar una action que realmente exista y especificar una versión válida:"
    )
    add_code_block(doc,
        'fix-dependencia-corregida:\n'
        '  runs-on: ubuntu-latest\n'
        '  steps:\n'
        '    - name: Usar action correcta\n'
        '      uses: actions/checkout@v4  # ← Action válida con versión correcta'
    )
    add_note_box(doc,
        "Lección aprendida: Siempre verificar que las actions referenciadas existan "
        "en el marketplace de GitHub y que la versión (tag) sea correcta. Usar el "
        "marketplace oficial para buscar actions verificadas.",
        "warning"
    )

    # 5.3
    doc.add_heading("5.3 Error: Path inválido", level=2)
    doc.add_heading("Descripción del error", level=3)
    doc.add_paragraph(
        "Se intenta ejecutar un script (./scripts/deploy_production.sh) y leer un "
        "archivo de configuración (/ruta/que/no/existe/config.yml) que no existen "
        "en el repositorio ni en el sistema de archivos."
    )
    add_code_block(doc,
        'error-path-invalido:\n'
        '  runs-on: ubuntu-latest\n'
        '  steps:\n'
        '    - uses: actions/checkout@v4\n'
        '    - name: Intentar ejecutar script inexistente\n'
        '      run: bash ./scripts/deploy_production.sh  # ← No existe\n'
        '    - name: Intentar acceder a archivo inexistente\n'
        '      run: cat /ruta/que/no/existe/config.yml  # ← No existe'
    )

    add_image_with_caption(doc, "error_path",
        "Job 'Error: Path Inválido' fallido con marca roja. Los errores muestran: "
        "'bash: ./scripts/deploy_production.sh: No such file or directory' para el "
        "script inexistente, y 'cat: /ruta/que/no/existe/config.yml: No such file or "
        "directory' para el archivo de configuración. Ambos errores son causados por "
        "intentar acceder a rutas que no existen en el sistema de archivos del runner."
    )

    doc.add_heading("Resolución", level=3)
    doc.add_paragraph(
        "La solución es crear los archivos necesarios antes de intentar acceder a ellos, "
        "o verificar su existencia con comprobaciones previas:"
    )
    add_code_block(doc,
        'fix-path-corregido:\n'
        '  steps:\n'
        '    - uses: actions/checkout@v4\n'
        '    - name: Crear y ejecutar script correctamente\n'
        '      run: |\n'
        '        mkdir -p scripts\n'
        '        cat > scripts/deploy_production.sh << \'EOF\'\n'
        '        #!/bin/bash\n'
        '        echo "🚀 Script ejecutándose correctamente"\n'
        '        EOF\n'
        '        chmod +x scripts/deploy_production.sh\n'
        '        bash ./scripts/deploy_production.sh'
    )
    add_note_box(doc,
        "Lección aprendida: Verificar siempre que los archivos y rutas referenciados "
        "en los workflows existan en el repositorio. Usar comprobaciones con 'test -f' "
        "o 'if [ -f archivo ]' antes de intentar ejecutar scripts.",
        "warning"
    )

    # 5.4
    doc.add_heading("5.4 Error: Permisos", level=2)
    doc.add_heading("Descripción del error", level=3)
    doc.add_paragraph(
        "Se intenta escribir en directorios del sistema (/etc/) y modificar permisos "
        "de archivos protegidos (/etc/passwd), operaciones que requieren privilegios "
        "de root que el runner no posee."
    )
    add_code_block(doc,
        'error-permisos:\n'
        '  steps:\n'
        '    - name: Escribir en directorio protegido\n'
        '      run: echo "test" > /etc/mi_config_app.conf  # ← Sin permisos\n'
        '    - name: Cambiar permisos de sistema\n'
        '      run: chmod 777 /etc/passwd  # ← Operación no permitida'
    )

    add_image_with_caption(doc, "error_permisos",
        "Job 'Error: Permisos' fallido con marca roja. Los errores muestran: "
        "'bash: /etc/mi_config_app.conf: Permission denied' al intentar escribir en /etc/, "
        "y 'chmod: changing permissions of /etc/passwd: Operation not permitted' al intentar "
        "modificar permisos de archivos del sistema. El runner de GitHub Actions ejecuta "
        "con permisos de usuario estándar, no como root."
    )

    doc.add_heading("Resolución", level=3)
    doc.add_paragraph(
        "La solución es trabajar únicamente dentro del directorio del workspace "
        "donde el runner tiene permisos de escritura:"
    )
    add_code_block(doc,
        'fix-permisos-corregido:\n'
        '  steps:\n'
        '    - name: Escribir en directorio con permisos\n'
        '      run: |\n'
        '        mkdir -p ${{ github.workspace }}/config\n'
        '        echo "config=true" > ${{ github.workspace }}/config/app.conf\n'
        '    - name: Gestionar permisos de archivos propios\n'
        '      run: |\n'
        '        touch mi_archivo.txt\n'
        '        chmod 644 mi_archivo.txt  # ← Solo archivos propios'
    )
    add_note_box(doc,
        "Lección aprendida: Los runners de GitHub Actions ejecutan con permisos "
        "limitados. Siempre trabajar dentro de $GITHUB_WORKSPACE y nunca intentar "
        "modificar archivos del sistema operativo.",
        "warning"
    )

    # 5.5
    doc.add_heading("5.5 Versiones corregidas", level=2)
    doc.add_paragraph(
        "Tras aplicar todas las correcciones, se ejecuta el workflow "
        "troubleshooting-fixed.yml que contiene las versiones corregidas de "
        "todos los errores anteriores:"
    )

    add_image_with_caption(doc, "fixed_success",
        "Ejecución exitosa del workflow 'Lab5 - Troubleshooting (Errores Corregidos)'. "
        "Los cuatro jobs muestran marcas verdes de éxito: 'Fix: Label Corregido', "
        "'Fix: Dependencia Corregida', 'Fix: Path Corregido' y 'Fix: Permisos Corregidos'. "
        "Todas las correcciones aplicadas funcionan correctamente, demostrando la "
        "resolución exitosa de cada problema identificado."
    )

    doc.add_page_break()

    # ==========================================
    # 6. PARTE 4 — LOGGING DETALLADO
    # ==========================================
    doc.add_heading("6. Parte 4 — Logging Detallado", level=1)

    # 6.1
    doc.add_heading("6.1 Configuración del logging verbose", level=2)
    doc.add_paragraph(
        "Para activar el logging detallado en GitHub Actions, se configuran las "
        "siguientes variables de entorno en el workflow:"
    )
    add_code_block(doc,
        'env:\n'
        '  ACTIONS_RUNNER_DEBUG: true   # Activa logs de depuración del runner\n'
        '  ACTIONS_STEP_DEBUG: true     # Activa logs de depuración por step'
    )
    doc.add_paragraph(
        "ACTIONS_RUNNER_DEBUG habilita logs adicionales del proceso del runner, "
        "incluyendo información sobre la conexión con GitHub, la descarga de actions "
        "y la gestión de la cola de jobs. ACTIONS_STEP_DEBUG proporciona información "
        "detallada sobre la ejecución de cada step individual."
    )
    doc.add_paragraph(
        "Estas variables también se pueden configurar como secrets del repositorio "
        "en Settings → Secrets → Actions para que apliquen a todos los workflows "
        "sin necesidad de modificar los archivos YAML."
    )

    # 6.2
    doc.add_heading("6.2 Análisis de los logs", level=2)
    doc.add_paragraph(
        "El workflow logging-detallado.yml ejecuta un diagnóstico completo en el "
        "runner self-hosted, recopilando información del sistema, conectividad y "
        "recursos disponibles."
    )

    add_image_with_caption(doc, "logging_verbose",
        "Salida del workflow de logging detallado mostrando la información de diagnóstico. "
        "Se observan las líneas con prefijo '##[debug]' que indican el logging verbose activado. "
        "La salida incluye variables del runner (RUNNER_NAME, RUNNER_OS: Windows, RUNNER_ARCH: X64), "
        "variables de GitHub (GITHUB_WORKFLOW, GITHUB_RUN_ID), información del sistema "
        "(Hostname: DESKTOP-A6A47AH) y confirmación de que ACTIONS_RUNNER_DEBUG y "
        "ACTIONS_STEP_DEBUG están activos. El logging detallado permite identificar "
        "exactamente dónde y por qué falla un workflow."
    )

    doc.add_paragraph(
        "El análisis de los logs verbose revela información crucial para el diagnóstico:"
    )
    diagnosis_items = [
        "Dónde falla: Los logs con prefijo ##[debug] muestran el punto exacto de fallo, incluyendo la línea de código, el step y la acción específica que causa el error.",
        "Qué contexto ayuda: Las variables de entorno del runner (RUNNER_NAME, RUNNER_OS, RUNNER_ARCH) permiten identificar en qué máquina se ejecuta el job y sus capacidades.",
        "Qué información da el runner: El runner reporta su versión, estado de conexión con GitHub, capacidad de resolución DNS, acceso a la API de GitHub, y métricas de recursos del sistema (CPU, memoria, disco).",
    ]
    for item in diagnosis_items:
        doc.add_paragraph(item, style='List Bullet')

    # 6.3
    doc.add_heading("6.3 Información del runner", level=2)
    doc.add_paragraph(
        "El workflow de diagnóstico recopila la siguiente información del runner:"
    )

    # Table
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = "Parámetro"
    headers[1].text = "Valor"
    for cell in headers:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data = [
        ("Runner Name", "LAB5-Runner"),
        ("Runner OS", "Windows"),
        ("Runner Architecture", "X64"),
        ("Computer Name", "DESKTOP-A6A47AH"),
        ("PowerShell Version", "5.1"),
        ("Git Version", "2.x"),
        ("Conectividad GitHub API", "✅ Accesible (Status 200)"),
    ]
    for i, (param, value) in enumerate(data):
        row = table.rows[i + 1].cells
        row[0].text = param
        row[1].text = value

    doc.add_paragraph()
    doc.add_page_break()

    # ==========================================
    # 7. PARTE 5 — SEGURIDAD
    # ==========================================
    doc.add_heading("7. Parte 5 — Seguridad", level=1)
    doc.add_paragraph(
        "Esta sección analiza los riesgos de seguridad asociados al uso de runners "
        "self-hosted en entornos de producción. Es fundamental comprender estos riesgos "
        "para implementar las medidas de mitigación adecuadas."
    )

    # 7.1
    doc.add_heading("7.1 Riesgos de runners persistentes", level=2)
    doc.add_paragraph(
        "A diferencia de los runners hosted que se destruyen después de cada ejecución, "
        "los runners self-hosted son persistentes, lo que introduce los siguientes riesgos:"
    )
    
    risks_persistent = [
        ("Datos residuales entre ejecuciones", 
         "Los archivos generados por un workflow (artefactos de build, credenciales "
         "temporales, datos de test) pueden permanecer en el disco del runner y ser "
         "accesibles por ejecuciones posteriores de otros repositorios o branches. "
         "Un atacante podría diseñar un workflow malicioso que busque y exfiltre estos datos."),
        ("Secretos en disco", 
         "Los secretos de GitHub Actions se inyectan como variables de entorno durante "
         "la ejecución. Si un workflow escribe estos secretos en archivos temporales "
         "o logs, podrían quedar persistidos en el disco del runner. Esto incluye tokens "
         "de API, contraseñas de bases de datos y claves de servicios cloud."),
        ("Malware persistente", 
         "Un workflow comprometido podría instalar software malicioso (backdoors, "
         "keyloggers, miners de criptomonedas) en el runner. Al ser persistente, este "
         "malware sobreviviría entre ejecuciones y podría afectar a todos los workflows "
         "futuros ejecutados en ese runner."),
        ("Modificación del entorno", 
         "Los workflows pueden modificar la configuración del sistema operativo, "
         "instalar dependencias o alterar variables de entorno de forma permanente, "
         "afectando la reproducibilidad y seguridad de ejecuciones futuras."),
    ]
    for title, desc in risks_persistent:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
        p.add_run(desc)
        doc.add_paragraph()

    doc.add_paragraph("Medidas de mitigación:")
    mitigations_1 = [
        "Ejecutar el runner en un contenedor o VM efímera que se recree después de cada job",
        "Implementar scripts de limpieza que se ejecuten al finalizar cada job",
        "Configurar el runner para que solo ejecute workflows de branches protegidos",
        "Monitorizar el espacio en disco y los procesos del runner regularmente",
    ]
    for m in mitigations_1:
        doc.add_paragraph(m, style='List Bullet')

    # 7.2
    doc.add_heading("7.2 Riesgos de acceso a red", level=2)
    doc.add_paragraph(
        "Los runners self-hosted se ejecutan dentro de la red corporativa, lo que "
        "les otorga acceso a recursos internos que no deberían estar expuestos:"
    )

    risks_network = [
        ("Acceso a red interna", 
         "El runner tiene acceso directo a la red corporativa, incluyendo servidores "
         "internos, bases de datos, APIs privadas y otros recursos que normalmente "
         "solo son accesibles desde la intranet. Un workflow malicioso podría escanear "
         "la red y descubrir servicios vulnerables."),
        ("Exfiltración de datos", 
         "Un atacante que comprometa un workflow podría utilizar el runner como punto "
         "de salida para exfiltrar datos corporativos sensibles. El runner tiene acceso "
         "a Internet (necesario para conectar con GitHub) y a la red interna, actuando "
         "como puente entre ambas redes."),
        ("Lateral movement", 
         "Si el runner se ejecuta con credenciales de dominio o tiene acceso a servicios "
         "de directorio activo, un atacante podría usar el runner como punto de pivote "
         "para moverse lateralmente por la red corporativa, comprometiendo otros sistemas."),
        ("DNS rebinding y SSRF", 
         "Los workflows pueden realizar peticiones HTTP a cualquier dirección accesible "
         "desde el runner. Esto incluye servicios internos, metadatos de cloud providers "
         "(169.254.169.254) y endpoints de administración que podrían revelar información "
         "sensible."),
    ]
    for title, desc in risks_network:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
        p.add_run(desc)
        doc.add_paragraph()

    doc.add_paragraph("Medidas de mitigación:")
    mitigations_2 = [
        "Aislar el runner en una VLAN dedicada con acceso restringido a la red interna",
        "Implementar firewall rules que limiten las conexiones salientes del runner",
        "Usar un proxy corporativo para filtrar y auditar el tráfico de red del runner",
        "No ejecutar el runner con credenciales de dominio con privilegios elevados",
        "Bloquear el acceso a endpoints de metadatos cloud (169.254.169.254)",
    ]
    for m in mitigations_2:
        doc.add_paragraph(m, style='List Bullet')

    # 7.3
    doc.add_heading("7.3 Riesgos de aislamiento insuficiente", level=2)
    doc.add_paragraph(
        "Los runners self-hosted que se ejecutan directamente en el sistema operativo "
        "(sin contenedores) presentan graves riesgos de aislamiento:"
    )

    risks_isolation = [
        ("Sin aislamiento de procesos", 
         "Los workflows se ejecutan como procesos del sistema operativo sin aislamiento. "
         "Esto significa que un workflow puede ver y potencialmente interactuar con otros "
         "procesos del sistema, incluyendo otros jobs que se ejecuten en paralelo y "
         "servicios del sistema operativo."),
        ("Acceso al sistema de archivos completo", 
         "Sin contenedores, el workflow tiene acceso al sistema de archivos completo "
         "del runner (limitado solo por permisos del usuario). Esto incluye archivos "
         "de configuración del sistema, otros repositorios, y archivos personales "
         "del usuario que ejecuta el runner."),
        ("Escalada de privilegios", 
         "Si el runner se ejecuta con permisos de administrador (o como un servicio "
         "del sistema), un workflow malicioso podría explotar vulnerabilidades del "
         "sistema operativo para obtener acceso root/SYSTEM, comprometiendo "
         "completamente la máquina."),
        ("Compartición de recursos", 
         "Múltiples jobs comparten CPU, memoria y disco sin límites. Un job malicioso "
         "podría consumir todos los recursos del sistema (crypto mining, fork bombs), "
         "afectando al rendimiento de otros jobs y del sistema host."),
    ]
    for title, desc in risks_isolation:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
        p.add_run(desc)
        doc.add_paragraph()

    doc.add_paragraph("Medidas de mitigación:")
    mitigations_3 = [
        "Ejecutar el runner dentro de un contenedor Docker para aislar cada ejecución",
        "Usar máquinas virtuales efímeras que se destruyan después de cada job",
        "Ejecutar el runner con una cuenta de usuario dedicada con privilegios mínimos",
        "Configurar limits de CPU, memoria y disco para los procesos del runner",
        "No ejecutar el runner como administrador o servicio del sistema con privilegios elevados",
        "Restringir los repositorios que pueden ejecutar jobs en el runner",
    ]
    for m in mitigations_3:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_page_break()

    # ==========================================
    # 8. CONCLUSIONES
    # ==========================================
    doc.add_heading("8. Conclusiones", level=1)
    doc.add_paragraph(
        "El Laboratorio 5 ha permitido practicar de forma integral la configuración, "
        "uso y troubleshooting de runners self-hosted en GitHub Actions. Las principales "
        "conclusiones son:"
    )

    conclusions = [
        "La configuración de runners self-hosted es un proceso relativamente sencillo pero "
        "que requiere atención a los detalles, especialmente en la asignación de labels "
        "para dirigir correctamente los jobs.",
        
        "Los pipelines híbridos (hosted + self-hosted) permiten combinar la conveniencia "
        "y escalabilidad de los runners de GitHub con la capacidad de ejecutar jobs en "
        "infraestructura propia cuando es necesario acceder a recursos internos.",
        
        "El troubleshooting en GitHub Actions requiere un enfoque metódico: identificar "
        "el tipo de error (label, dependencia, path, permisos), analizar los logs y "
        "aplicar la corrección adecuada. Los errores más comunes son fáciles de resolver "
        "una vez identificados.",
        
        "El logging detallado (ACTIONS_RUNNER_DEBUG y ACTIONS_STEP_DEBUG) es una herramienta "
        "fundamental para el diagnóstico de problemas. Proporciona información invaluable "
        "sobre el estado del runner, las variables de entorno y el flujo de ejecución.",
        
        "La seguridad de los runners self-hosted es un aspecto crítico que no debe subestimarse. "
        "Los riesgos de persistencia, acceso a red y falta de aislamiento requieren implementar "
        "medidas de mitigación robustas, especialmente en entornos corporativos.",
    ]
    for i, conclusion in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.font.bold = True
        p.add_run(conclusion)
        doc.add_paragraph()

    doc.add_heading("Entregables completados", level=2)
    deliverables = [
        ("✅ Runner funcionando", "Runner LAB5-Runner registrado y operativo con labels self-hosted, windows, lab5"),
        ("✅ Pipeline híbrido", "Workflow hybrid-pipeline.yml con 3 jobs (2 hosted + 1 self-hosted)"),
        ("✅ Documento de troubleshooting", "5 errores documentados con causa raíz, capturas y resolución"),
        ("✅ Explicación de riesgos seguridad", "Análisis detallado de 3 categorías de riesgos con medidas de mitigación"),
    ]
    for title, desc in deliverables:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        p.add_run(desc)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"[OK] Documento guardado en: {OUTPUT_FILE}")
    print(f"     Tamano: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
